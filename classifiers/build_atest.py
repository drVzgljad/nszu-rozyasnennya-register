# -*- coding: utf-8 -*-
"""
Номенклатура атестації — розбір Додатка 1 до Порядку проведення атестації
працівників сфери охорони здоров'я (наказ МОЗ від 16.04.2025 № 650, z0824-25).

НАВІЩО. Додаток 1 — «Номенклатура спеціальностей / профілів роботи за
спеціальностями та відповідних їм професійних кваліфікацій / посад» — ЄДИНЕ
офіційне зіставлення «спеціальність ↔ кваліфікація ↔ посада» в кадровому
контурі. Решта наших містків (Додаток 7 ↔ Перелік 1065, Перелік ↔ ДКХП)
обчислені за назвами, бо їх не існує в актах. Тут акт зіставляє сам.

ЩО ЧИТАЄ
--------
  0_10_класификатори/спеціальності/z0824-25-dod1.docx
      ← сигнальний документ Додатка 1 (стягує fetch_sources.py).
        У /print наказу 650 самих таблиць немає — лише плашки-підписи
        додатків; та сама пастка, що з Додатком 7 ПКМУ 813.

ДВІ ТАБЛИЦІ ДОДАТКА
-------------------
  1. «Спеціальності та професійні кваліфікації, що здобуваються на їх
     основі» — атестаційна спеціальність → кваліфікація(-ї). Кваліфікацій у
     клітинці буває кілька (розділені новим рядком): «Судова медицина» дає
     і лаборанта, і фельдшера-лаборанта.
  2. «Профілі роботи за спеціальностями та відповідні посади/функції» —
     профіль → посада/функція після субспеціалізації → базові спеціальності.
     Саме звідси родом 63 «профільні» позиції Переліку МОЗ № 1065.

ПАСТКИ
------
  1. Секційні рядки — merged cells: python-docx повторює текст у КОЖНІЙ
     клітинці рядка. Ознака секції — всі клітинки однакові (uniq == 1).
  2. Три рівні секцій: «І.» (римська, кирилична І!), «1.» (крапка),
     «1)» (дужка). Нумерація данських рядків починається наново в кожній
     секції — рядок «1 2 3 4» з номерами колонок теж не даний (та сама
     пастка, що коштувала 39 позицій у Переліку 1065).
  3. Зірочки в назвах спеціальностей (*, **, ***) — виноски під таблицею.
     З назви їх знімаємо, кількість лишаємо в полі stars, текст виносок —
     у meta: «Зубний лікар» припинено готувати саме у виносці ***.
  4. Переноси в клітинках — явні \n; склеюємо пробілом. Дослівні дефекти
     джерела («Лікар ревматолог дитячий» без дефіса) НЕ виправляємо.
  5. Колонка «віднесення… до переліку спеціальностей МОН» — «+» або
     порожньо; це прапорець освітнього переліку, не наш місток.
  6. У профілів колонка спеціальностей буває НЕ списком, а текстом
     («Лікарські спеціальності, крім спеціальностей стоматологічного
     профілю») — такий рядок лишаємо одним елементом, зіставляти його з
     конкретними спеціальностями не можна.

ВИХІД: classifiers/data/atest/
  atest_specs.json     [{id, name, path, mon, quals, stars}]
  atest_profiles.json  [{id, name, path, mon, posts_text, base_specs}]
  atest_meta.json      реквізити акта, виноски, звірочні числа

Код повернення: 0 — числа збіглися з контрольними; 2 — розійшлися
(джерело оновилося). Дані пишуться в обох випадках.
"""
import json
import re
import sys
from datetime import date
from pathlib import Path

import docx

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

BASE = Path(__file__).resolve().parent
SRC = BASE.parent.parent / "0_10_класификатори" / "спеціальності"
DOCX = SRC / "z0824-25-dod1.docx"
OUT_DIR = BASE / "data" / "atest"

# Контрольні числа — знімок Додатка 1 у редакції від 26.06.2026 (зміни
# наказами № 618 від 14.05.2026 і № 754 від 04.06.2026). Розійшлися —
# отже, МОЗ оновив Номенклатуру, і правити словник треба свідомо.
# 192 = 129 лікарських + 7 фармацевтичних + 28 професіоналів + 28 фахівців.
EXPECTED = {
    "specs": 192,
    "profiles": 51,
    "quals": 197,
    "sections_t1": 4,
}

ROMAN_RE = re.compile(r"^[IVХІ]+\.\s")      # римські: латинські І кириличні
SUB_RE = re.compile(r"^\d+\.\s")
SUBSUB_RE = re.compile(r"^\d+\)\s")


def clean(s):
    """Клітинка → рядок: явні переноси і nbsp — у пробіли, дослівність назв
    не чіпаємо."""
    return re.sub(r"\s+", " ", s.replace("\xa0", " ")).strip()


def lines(s):
    """Клітинка → список рядків (кваліфікації, базові спеціальності)."""
    out = []
    for part in s.replace("\xa0", " ").split("\n"):
        part = re.sub(r"\s+", " ", part).strip()
        if part:
            out.append(part)
    # Перенос усередині ОДНІЄЇ назви теж дає \n («Загальна практика — \n
    # сімейна медицина»). Відрізняємо від переліку так: хвіст із малої
    # літери, після тире, або дужка парної форми («Сестра медична-анестезист
    # \n (брат медичний-анестезист)») — це продовження попереднього рядка.
    merged = []
    for part in out:
        if merged and (part[:1].islower() or part[:1] == "("
                       or merged[-1].endswith(("—", "-", "–"))):
            merged[-1] = merged[-1] + " " + part
        else:
            merged.append(part)
    return merged


def stars_of(name):
    m = re.search(r"(\*+)\s*$", name)
    return (re.sub(r"\*+\s*$", "", name).strip(), len(m.group(1)) if m else 0)


def parse_table(table, id_prefix, data_cols):
    """Спільний прохід: секції за uniq==1, номерний рядок і шапку геть."""
    path = {"sec": None, "sub": None, "subsub": None}
    rows = []
    for r in table.rows:
        cells = [c.text for c in r.cells]
        flat = [clean(c) for c in cells]
        uniq = {c for c in flat if c}
        if not uniq:
            continue
        if len(uniq) == 1:                       # секційний рядок (merged)
            t = uniq.pop()
            if ROMAN_RE.match(t):
                path = {"sec": t, "sub": None, "subsub": None}
            elif SUB_RE.match(t):
                path["sub"], path["subsub"] = t, None
            elif SUBSUB_RE.match(t):
                path["subsub"] = t
            continue
        # шапка і рядок нумерації колонок
        if flat[0].startswith("№") or all(re.fullmatch(r"\d", c) for c in flat if c):
            continue
        if not re.fullmatch(r"\d+", flat[0] or ""):
            continue
        rows.append((dict(path), cells))
    out = []
    for i, (p, cells) in enumerate(rows, 1):
        rec = {"id": f"{id_prefix}{i:03d}",
               "path": [x for x in (p["sec"], p["sub"], p["subsub"]) if x]}
        rec.update(data_cols(cells))
        out.append(rec)
    return out


def spec_cols(cells):
    name, stars = stars_of(clean(cells[1]))
    return {"name": name, "stars": stars,
            "mon": clean(cells[2]) == "+",
            "quals": [stars_of(q)[0] for q in lines(cells[3])]}


def profile_cols(cells):
    name, stars = stars_of(clean(cells[1]))
    # Посад у клітинці буває кілька (окремими рядками): «Лаборант з
    # імунології» + «Фельдшер-лаборант з імунології». Склеїти їх в один
    # рядок означало б загубити обидві при зіставленні з Переліком.
    return {"name": name, "stars": stars,
            "mon": clean(cells[2]) == "+",
            "posts": lines(cells[3]),
            "base_specs": lines(cells[4])}


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    d = docx.Document(str(DOCX))
    if len(d.tables) != 2:
        sys.exit(f"У Додатку 1 очікуємо 2 таблиці, знайдено {len(d.tables)}")

    specs = parse_table(d.tables[0], "AS", spec_cols)
    profiles = parse_table(d.tables[1], "AP", profile_cols)

    footnotes = [p.text.strip() for p in d.paragraphs
                 if p.text.strip().startswith("*")]

    sources = json.loads((SRC / "sources.json").read_text(encoding="utf-8"))
    rec650 = next((x for x in sources["documents"] if x["key"] == "z0824-25"), {})

    actual = {
        "specs": len(specs),
        "profiles": len(profiles),
        "quals": sum(len(s["quals"]) for s in specs),
        "sections_t1": len({p["path"][0] for p in specs if p["path"]}),
    }
    ok = actual == EXPECTED

    def write(name, obj):
        p = OUT_DIR / name
        p.write_text(json.dumps(obj, ensure_ascii=False, separators=(",", ":")),
                     encoding="utf-8")
        return p

    write("atest_specs.json", specs)
    write("atest_profiles.json", profiles)
    write("atest_meta.json", {
        "generated": date.today().isoformat(),
        "act": {
            "title": "наказ МОЗ від 16.04.2025 № 650",
            "reg": "Мін'юст 28.05.2025 № 824/44230",
            "rada": "z0824-25",
            "status": rec650.get("status", ""),
            "revision": rec650.get("revision", ""),
            "note_war": ("Процедура атестації запускається через 6 місяців "
                         "після припинення чи скасування воєнного стану; строк "
                         "дії сертифікатів і посвідчень, що сплив у період "
                         "воєнного стану, подовжено"),
        },
        # Числа БПР — з розділу VIII Порядку і прикінцевих положень наказу
        # (звірено з /print z0824-25 12.08.2026): мінімум за рік, сумарний
        # поріг за атестаційний період з 2029 року і перехідна драбина
        # 2026–2028 на випадок відновлення атестації в ці роки.
        "bpr": {
            "period_max_years": 5,
            "professional": {"per_year": 50, "period_from_2029": 250,
                             "transition": {"2026": 100, "2027": 150, "2028": 200}},
            "fakhivets": {"per_year": 30, "period_from_2029": 150,
                          "transition": {"2026": 60, "2027": 90, "2028": 120}},
        },
        "footnotes": footnotes,
        "check": {"ok": ok, "expected": EXPECTED, "actual": actual},
    })

    print(f"Спеціальностей: {len(specs)} · кваліфікацій: {actual['quals']} · "
          f"профілів: {len(profiles)}")
    for p in specs[:3] + specs[-2:]:
        print("  S:", p["id"], p["name"], "→", "; ".join(p["quals"])[:60])
    for p in profiles[:2]:
        print("  P:", p["id"], p["name"], "→", "; ".join(p["posts"])[:50])
    print("Виносок:", len(footnotes))
    print("Звірка:", "OK" if ok else f"РОЗІЙШЛОСЯ {actual} != {EXPECTED}")
    return 0 if ok else 2


if __name__ == "__main__":
    sys.exit(main())
