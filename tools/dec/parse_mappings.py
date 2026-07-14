import os
import re
import json
import docx
import collections
from pathlib import Path

# Paths
BASE_DIR = Path(__file__).resolve().parent
DOCX_PATH = BASE_DIR / ".." / ".." / "0_10_класификатори" / "Таблиця_співставлення_2026_03__ (1).docx"
PACKAGES_JSON_PATH = BASE_DIR / ".." / "pakety" / "data" / "packages_2026.json"
OUTPUT_JSON_PATH = BASE_DIR / "data" / "package_codes.json"

def main():
    print(f"Loading mapping document from {DOCX_PATH}...")
    if not DOCX_PATH.exists():
        print(f"Error: docx file not found at {DOCX_PATH}")
        return
        
    doc = docx.Document(DOCX_PATH)
    
    # 1. Parse Table 2 (DRG / ОДК)
    odk_map = {}
    t2 = doc.tables[1]
    for row in t2.rows[1:]:
        odk_code = row.cells[0].text.strip() # e.g. 'ОДК 1'
        odk_codes_str = row.cells[2].text.strip()
        
        # Split codes by comma, space, semicolon
        raw_codes = re.split(r'[\s,;]+', odk_codes_str)
        clean_codes = [c.strip() for c in raw_codes if c.strip() and c.strip() != '-']
        odk_map[odk_code] = clean_codes
        
    print(f"Parsed {len(odk_map)} ODK maps from Table 2.")
    
    # 2. Load Package names from packages_2026.json
    package_names = {}
    if PACKAGES_JSON_PATH.exists():
        try:
            with open(PACKAGES_JSON_PATH, "r", encoding="utf-8") as f:
                pkg_data = json.load(f)
                for p in pkg_data.get("packages", []):
                    package_names[p["number"]] = p["title"]
            print(f"Loaded {len(package_names)} package names from packages_2026.json.")
        except Exception as e:
            print(f"Warning: Failed to load packages_2026.json: {e}")
            
    # 3. Parse Table 1 (Mappings)
    package_data = collections.defaultdict(lambda: {'icd10_targets': set(), 'achi_targets': set()})
    t1 = doc.tables[0]
    
    for r_idx, row in enumerate(t1.rows[1:], 1):
        pkg_str = row.cells[0].text.strip()
        icd_str = row.cells[2].text.strip()
        achi_str = row.cells[3].text.strip()
        
        # Parse packages (can be '3', '3, 47')
        pkg_nums = [p.strip() for p in re.split(r'[\s,;]+', pkg_str) if p.strip()]
        if not pkg_nums:
            continue
            
        # Parse ICD-10
        icd_codes = []
        if icd_str and icd_str != '-':
            # Check if it references ODK
            odk_refs = re.findall(r'ОДК\s*\d+', icd_str)
            if odk_refs:
                for ref in odk_refs:
                    norm_ref = re.sub(r'\s+', ' ', ref)
                    
                    # Check exact match
                    if norm_ref in odk_map:
                        icd_codes.extend(odk_map[norm_ref])
                    else:
                        # Try prefix match (e.g. ОДК 21 matches ОДК 21A, ОДК 21B)
                        matched = False
                        for key in odk_map:
                            if key.startswith(norm_ref) and (len(key) == len(norm_ref) or not key[len(norm_ref)].isdigit()):
                                icd_codes.extend(odk_map[key])
                                matched = True
                        if not matched:
                            print(f"Warning: ODK ref {ref} at row {r_idx} not resolved")
            else:
                raw_icd = re.split(r'[\s,;]+', icd_str)
                icd_codes.extend([c.strip() for c in raw_icd if c.strip() and c.strip() != '-'])
                
        # Parse ACHI
        achi_codes = []
        if achi_str and achi_str != '-':
            raw_achi = re.split(r'[\s,;]+', achi_str)
            achi_codes.extend([c.strip() for c in raw_achi if c.strip() and c.strip() != '-'])
            
        for p in pkg_nums:
            package_data[p]['icd10_targets'].update(icd_codes)
            package_data[p]['achi_targets'].update(achi_codes)
            
    # 4. Format the final output dictionary
    output_data = {}
    for p in sorted(package_data.keys(), key=lambda x: int(x) if x.isdigit() else 999):
        # Resolve package name
        name = package_names.get(p, f"Пакет {p}")
        
        # Sort targets for clean output
        icd_sorted = sorted(list(package_data[p]['icd10_targets']))
        achi_sorted = sorted(list(package_data[p]['achi_targets']))
        
        output_data[p] = {
            "package_name": name,
            "icd10_targets": icd_sorted,
            "achi_targets": achi_sorted
        }
        
    # 5. Write to package_codes.json
    OUTPUT_JSON_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_JSON_PATH, "w", encoding="utf-8") as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
        
    print(f"Successfully compiled mapping database: {len(output_data)} packages written to {OUTPUT_JSON_PATH} ({OUTPUT_JSON_PATH.stat().st_size} bytes)")

if __name__ == "__main__":
    main()
