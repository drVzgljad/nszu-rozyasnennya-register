import json
import re
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


BASE = Path(__file__).resolve().parents[1]
REGISTRY_PATH = BASE / "04_Реєстр" / "реєстр_розяснень_НСЗУ_175.json"
ORIGINALS = BASE / "00_Оригінали"
TEXT_DIR = BASE / "01_Текст_та_OCR"
OUTPUT_PATH = BASE / "04_Реєстр" / "аналіз_текстового_шару.json"


def normalize(text):
    return re.sub(r"\s+", " ", text or "").strip()


def save_text(stem, content):
    target = TEXT_DIR / f"{stem}.txt"
    target.write_text(content, encoding="utf-8")
    return str(target.relative_to(BASE))


def extract_pdf(path):
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    text = "\n\n".join(pages)
    clean = normalize(text)
    chars = len(clean)
    page_count = len(reader.pages)
    per_page = chars / page_count if page_count else 0
    if chars < 80 or per_page < 20:
        quality = "needs_ocr"
    elif per_page < 150:
        quality = "text_sparse_review"
    else:
        quality = "text_ok"
    return text, {"pages": page_count, "text_chars": chars, "text_quality": quality}


def extract_docx(path):
    doc = Document(str(path))
    fragments = []
    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if text:
            fragments.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [normalize(cell.text) for cell in row.cells]
            if any(cells):
                fragments.append(" | ".join(cells))
    text = "\n".join(fragments)
    return text, {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "text_chars": len(normalize(text)),
        "text_quality": "text_ok" if normalize(text) else "empty_review",
    }


def extract_xlsx(path):
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    fragments = []
    sheets = []
    for sheet in workbook.worksheets:
        sheets.append(sheet.title)
        fragments.append(f"[Аркуш: {sheet.title}]")
        written = 0
        for row in sheet.iter_rows(values_only=True):
            cells = [normalize(str(value)) for value in row if value is not None]
            if cells:
                fragments.append(" | ".join(cells))
                written += 1
            if written >= 200:
                fragments.append("[...витяг обмежено 200 непорожніми рядками...]")
                break
    text = "\n".join(fragments)
    return text, {
        "sheets": sheets,
        "text_chars": len(normalize(text)),
        "text_quality": "text_ok" if normalize(text) else "empty_review",
    }


def extract_zip(path):
    with zipfile.ZipFile(path) as archive:
        entries = archive.namelist()
    text = "\n".join(entries)
    return text, {
        "zip_entries": len(entries),
        "text_chars": len(text),
        "text_quality": "archive_review",
    }


def main():
    registry = json.loads(REGISTRY_PATH.read_text(encoding="utf-8"))
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    by_file = {}
    for row in registry["documents"]:
        by_file.setdefault(row["original_file_name"], []).append(row)

    analyses = []
    for file_name, rows in by_file.items():
        path = ORIGINALS / file_name
        ext = path.suffix.lower().lstrip(".")
        item = {
            "original_file_name": file_name,
            "record_ids": [row["id"] for row in rows],
            "extension": ext,
            "site_titles": [row["site_title"] for row in rows],
            "exists": path.exists(),
        }
        try:
            if ext == "pdf":
                text, meta = extract_pdf(path)
            elif ext == "docx":
                text, meta = extract_docx(path)
            elif ext == "xlsx":
                text, meta = extract_xlsx(path)
            elif ext == "zip":
                text, meta = extract_zip(path)
            else:
                text, meta = "", {"text_quality": "unsupported_review", "text_chars": 0}
            item.update(meta)
            item["text_file"] = save_text(path.stem, text)
            item["text_preview"] = normalize(text)[:800]
        except Exception as exc:
            item.update({"text_quality": "extract_error", "text_chars": 0, "error": str(exc)})
        analyses.append(item)

    quality_counts = Counter(item["text_quality"] for item in analyses)
    extension_counts = Counter(item["extension"] for item in analyses)
    output = {
        "records": len(registry["documents"]),
        "unique_files": len(analyses),
        "duplicates": len(registry["documents"]) - len(analyses),
        "extension_counts": dict(extension_counts),
        "text_quality_counts": dict(quality_counts),
        "files": analyses,
    }
    OUTPUT_PATH.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({key: output[key] for key in ("records", "unique_files", "duplicates", "extension_counts", "text_quality_counts")}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
